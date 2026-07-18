from datetime import date
from pathlib import Path
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = Path(r'G:\AI-law-creater')
SRC = sorted((BASE / 'downloads').glob('*.xlsx'), key=lambda p: p.stat().st_mtime, reverse=True)[0]
OUT = BASE / 'AI_RAG知识库系统_需求文档.docx'

wb = load_workbook(SRC, data_only=True)


def cell_text(v):
    if v is None:
        return ''
    if hasattr(v, 'strftime'):
        return v.strftime('%Y-%m-%d')
    return str(v).strip()


def table_rows(sheet_name, header_row):
    ws = wb[sheet_name]
    headers = [cell_text(c.value) for c in ws[header_row] if cell_text(c.value)]
    out = []
    for row in ws.iter_rows(min_row=header_row+1, max_col=len(headers), values_only=True):
        vals = [cell_text(v) for v in row]
        if any(vals):
            out.append(dict(zip(headers, vals)))
    return headers, out

phase_headers, phases = table_rows('总体周期规划', 4)
wbs_headers, wbs = table_rows('WBS甘特计划', 4)
milestone_headers, milestones = table_rows('里程碑节点', 4)
risk_headers, risks = table_rows('风险登记册', 4)
quality_headers, qualities = table_rows('质量与验收标准', 4)
comm_headers, comms = table_rows('沟通计划', 4)
budget_headers, budgets = table_rows('预算与采购估算', 4)
raci_headers, racis = table_rows('资源与RACI', 4)

# Project overview values
ws_over = wb['项目总览']
overview = {}
for r in range(4, 9):
    vals = [cell_text(ws_over.cell(r, c).value) for c in range(1, 9)]
    for i in range(0, len(vals), 2):
        if vals[i]:
            overview[vals[i]] = vals[i+1] if i+1 < len(vals) else ''

# Derived requirement catalog
functional_requirements = [
    ('FR-001', '知识源管理', '系统应支持登记、分类、启停、版本标识和责任人维护知识源，覆盖首批核心业务文档。', '高', '知识源清单、入库日志、业务确认记录'),
    ('FR-002', '文档上传与接入', '系统应支持 Word、PDF、Excel、PPT、Markdown、网页文本等常见知识材料接入，并记录来源、更新时间、权限标签。', '高', '样例文件成功接入并可追溯来源'),
    ('FR-003', '文档解析与清洗', '系统应完成文本抽取、表格处理、OCR/图片文本识别、去重、噪声清洗和结构化元数据补全。', '高', '首批样例文档解析准确，异常文件进入待处理队列'),
    ('FR-004', '知识切分策略', '系统应支持按标题层级、段落、语义窗口、表格单元等策略切分，并保存片段与原文位置映射。', '高', '切分片段可定位到原文，切分策略可配置'),
    ('FR-005', 'Embedding 向量化', '系统应对知识片段生成向量，支持批量向量化、失败重试、增量更新与模型版本记录。', '高', '向量索引可用于检索，更新后版本可追踪'),
    ('FR-006', '向量检索', '系统应支持基于语义相似度的 Top-K 检索，并按知识库、权限、标签、时间、文档类型过滤。', '高', '黄金问答集 Top-5 召回命中率达到验收门槛'),
    ('FR-007', '混合检索与重排', '系统宜支持关键词检索、向量检索、元数据过滤、重排模型组合，以提升复杂问题命中率。', '中', '对比评测显示混合检索优于单一检索'),
    ('FR-008', 'RAG 问答编排', '系统应完成问题改写、检索、上下文组装、提示词模板、LLM 生成、引用输出的端到端流程。', '高', '端到端问答可演示且日志完整'),
    ('FR-009', '引用溯源', '系统回答应展示引用来源，包括文档标题、片段位置、更新时间和可打开的原文链接。', '高', '引用完整性达到 ≥95%'),
    ('FR-010', '低置信度拒答', '当检索结果不足、权限不足或置信度低时，系统应拒答或提示需要补充资料，不得编造答案。', '高', '无依据回答被拦截，拒答话术符合业务要求'),
    ('FR-011', '多轮会话', '系统应支持上下文相关追问，并可对会话上下文长度、保留轮次、敏感信息进行控制。', '中', '多轮问题保持语义连贯且不越权'),
    ('FR-012', '用户反馈闭环', '系统应支持点赞、点踩、纠错、补充知识建议、反馈分派和关闭状态跟踪。', '高', '反馈可记录、分类、分派、关闭'),
    ('FR-013', '管理后台', '系统应提供知识库、知识源、处理任务、模型配置、提示词模板、评测集、用户反馈、运行日志等管理能力。', '高', '管理员可完成核心运维流程'),
    ('FR-014', '权限控制', '系统应继承或配置文档级、片段级、知识库级访问权限，支持 RBAC/ABAC 或组织架构权限映射。', '高', '用户无法检索或生成无权限内容'),
    ('FR-015', '审计日志', '系统应记录用户查询、检索片段、模型调用、权限过滤、答案输出、反馈和管理员操作日志。', '高', '关键操作可审计、可追溯'),
    ('FR-016', '评测体系', '系统应维护黄金问答集，支持检索命中率、答案有用率、引用完整性、拒答准确性、时延和成本评测。', '高', 'MVP、UAT、上线前均形成评测报告'),
    ('FR-017', '运营看板', '系统应展示使用量、活跃用户、问题分类、命中率、反馈处理、Token 成本、响应时延、失败率等指标。', '中', '运营角色可按周期查看质量与成本趋势'),
    ('FR-018', '知识更新流程', '系统应支持新增、修改、删除、失效知识后的增量同步、重新向量化和索引刷新。', '高', '知识更新后可在约定时限内生效'),
    ('FR-019', '接口与集成', '系统应预留与 SSO、组织架构、企业网盘、IM/门户、监控告警、日志平台的集成接口。', '中', '完成至少核心身份与知识源集成'),
    ('FR-020', '上线与运维支持', '系统应提供部署配置、监控告警、备份恢复、回滚预案和运维手册。', '高', '上线评审和运维移交通过'),
]

non_functional = [
    ('NFR-001', '安全性', '用户不得访问无权限知识；高危漏洞为 0；敏感信息应脱敏或按策略拦截。', '高'),
    ('NFR-002', '性能', 'P95 首字响应 ≤ 3 秒，完整回答 ≤ 12 秒；可根据模型部署方式在基线评审时调整。', '高'),
    ('NFR-003', '可用性', '试运行期间核心服务可用性 ≥ 99.5%，重大故障为 0。', '高'),
    ('NFR-004', '可扩展性', '知识源、模型、向量库、重排器和提示词模板应具备可替换能力，避免强供应商锁定。', '中'),
    ('NFR-005', '可观测性', '应具备请求链路追踪、模型调用监控、检索质量监控、成本监控、告警与报表。', '高'),
    ('NFR-006', '可维护性', '配置、代码、部署脚本、接口文档、运维手册完整，支持交接和后续版本迭代。', '高'),
    ('NFR-007', '合规性', '数据接入、存储、处理、调用第三方模型应满足企业数据分级、隐私和审计要求。', '高'),
    ('NFR-008', '成本可控', '应提供 Token、模型调用、向量库和计算资源成本统计、预算告警及限流策略。', '中'),
]

roles = [
    ('普通用户', '提交自然语言问题、查看回答与引用、进行反馈。'),
    ('业务专家', '维护黄金问答集、参与知识质量审核、处理专业反馈、参与验收。'),
    ('知识运营', '管理知识源、跟踪反馈闭环、发布运营月报、组织知识更新。'),
    ('系统管理员', '维护用户权限、模型配置、提示词模板、任务队列、日志与告警。'),
    ('安全/合规人员', '审核数据分级、权限模型、审计日志、安全测试和上线准入。'),
    ('运维人员', '负责部署、监控、备份恢复、容量规划、上线与超保支持。'),
]

use_cases = [
    ('UC-001', '知识问答', '用户输入问题，系统基于权限过滤后的知识片段生成答案并展示引用。'),
    ('UC-002', '知识入库', '管理员选择知识源，系统解析、清洗、切分、向量化并发布到知识库。'),
    ('UC-003', '知识更新', '知识源发生变更，系统进行增量同步和索引刷新。'),
    ('UC-004', '答案纠错', '用户反馈答案错误，运营或业务专家复核后更新知识或调整提示词/评测集。'),
    ('UC-005', '质量评测', '测试或算法人员使用黄金问答集评估检索、答案、引用、拒答和性能指标。'),
    ('UC-006', '安全审计', '安全人员抽查用户查询、引用片段、权限过滤和管理员操作日志。'),
]

interfaces = [
    ('身份认证/SSO', '获取用户身份、组织架构、角色和权限属性。'),
    ('企业知识源', '对接网盘、文档库、业务系统、网页或数据库中的授权知识材料。'),
    ('模型服务', '调用 LLM、Embedding、重排模型，可支持公有云 API 或私有化模型。'),
    ('向量数据库', '存储向量索引、元数据、权限标签和检索配置。'),
    ('监控日志平台', '输出调用链路、错误、性能、成本和审计日志。'),
    ('消息/门户入口', '可将问答能力嵌入企业门户、IM 或工作台。'),
]

# Word helpers

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = str(h)
        set_cell_shading(cell, '1F4E78')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = cell_text(val)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run.font.size = Pt(9)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.color.rgb = RGBColor(31, 78, 121)
    return p


def add_para(doc, text='', bold_label=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_label:
        r = p.add_run(bold_label)
        r.bold = True
        r.font.name = '微软雅黑'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    r = p.add_run(text)
    r.font.name = '微软雅黑'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    r.font.size = Pt(10.5)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.0)

styles = doc.styles
styles['Normal'].font.name = '微软雅黑'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
styles['Normal'].font.size = Pt(10.5)

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(120)
r = p.add_run('AI RAG知识库系统\n需求规格说明书')
r.bold = True
r.font.size = Pt(24)
r.font.name = '微软雅黑'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
r.font.color.rgb = RGBColor(31, 78, 121)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('版本：V1.0\n编制日期：2026-07-16\n依据材料：飞书节点 AI_RAG知识库系统_项目计划表.xlsx')
r.font.size = Pt(12)
r.font.name = '微软雅黑'
r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

add_heading(doc, '1. 文档概述', 1)
add_para(doc, '本文档基于飞书节点中的《AI_RAG知识库系统_项目计划表.xlsx》编制，将项目计划中的范围、里程碑、WBS、质量验收、风险和治理要求转化为可用于需求评审、研发设计、测试验收和上线移交的需求规格说明。')
add_table(doc, ['项目项', '内容'], [
    ['项目名称', overview.get('项目名称', 'AI RAG 知识库系统建设项目')],
    ['项目周期', overview.get('项目周期', '2026-07-20 至 2027-01-15')],
    ['项目目标', overview.get('项目目标', '')],
    ['核心价值', overview.get('核心价值', '')],
    ['项目范围', overview.get('项目范围', '')],
    ['范围外', overview.get('范围外', '')],
    ['关键假设', overview.get('关键假设', '')],
    ['主要约束', overview.get('主要约束', '')],
], widths=[4, 13])

add_heading(doc, '2. 背景与建设目标', 1)
add_para(doc, '企业知识通常分散在网盘、文档、业务系统、邮件或人工经验中，存在查找成本高、口径不一致、知识更新慢、无法审计和难以复用等问题。本项目建设 AI RAG 知识库系统，通过检索增强生成技术，把可信知识片段注入大模型回答过程，形成可溯源、可运营、可审计的知识问答能力。')
add_para(doc, '建设目标包括：')
for item in [
    '完成核心知识源接入、解析、清洗、切分、向量化和索引管理。',
    '提供面向用户的自然语言问答能力，并输出明确引用来源。',
    '建立权限过滤、日志审计、低置信度拒答和反馈闭环，控制幻觉与越权风险。',
    '建立黄金问答集和质量评测体系，为 MVP、UAT 和上线验收提供量化依据。',
    '形成可持续运营的知识治理、质量监控、成本监控和版本迭代机制。',
]:
    add_para(doc, item, bold_label='• ')

add_heading(doc, '3. 项目范围', 1)
add_heading(doc, '3.1 范围内', 2)
for item in [
    '知识源接入、文档解析、清洗、切分、元数据管理与向量化。',
    'Embedding、向量数据库、检索、重排、RAG 编排和提示词模板管理。',
    '面向用户的问答界面、引用溯源、反馈提交和多轮会话能力。',
    '管理后台，包括知识库、知识源、处理任务、模型配置、评测集、反馈和日志管理。',
    '权限控制、数据分级、安全审计、合规检查、性能测试和上线运维移交。',
]:
    add_para(doc, item, bold_label='• ')
add_heading(doc, '3.2 范围外', 2)
for item in [
    '替代正式法律、医疗、财务等高风险场景的专业审查或最终决策。',
    '未授权数据采集、绕过原系统权限的数据处理。',
    '大规模业务系统重构或与 RAG 能力无直接关系的流程再造。',
    '未纳入当前阶段的全量多语言治理、复杂工作流自动执行和大规模多租户商业化能力。',
]:
    add_para(doc, item, bold_label='• ')

add_heading(doc, '4. 相关方与用户角色', 1)
add_table(doc, ['角色', '职责/使用场景'], roles, widths=[4, 12])

add_heading(doc, '5. 业务流程与核心用例', 1)
add_para(doc, '系统核心业务闭环为：知识接入 → 解析清洗 → 切分与元数据 → 向量化入库 → 权限过滤检索 → RAG 生成回答 → 引用溯源 → 用户反馈 → 运营复核 → 知识更新与评测优化。')
add_table(doc, ['用例编号', '用例名称', '说明'], use_cases, widths=[3, 4, 10])

add_heading(doc, '6. 功能需求', 1)
add_table(doc, ['编号', '模块', '需求描述', '优先级', '验收依据'], functional_requirements, widths=[2.2, 3, 8, 2, 4])

add_heading(doc, '7. 非功能需求', 1)
add_table(doc, ['编号', '维度', '需求描述', '优先级'], non_functional, widths=[2.2, 3, 10, 2])

add_heading(doc, '8. 数据与知识治理需求', 1)
for item in [
    '知识源应建立唯一标识、负责人、业务域、敏感级别、更新频率、接入方式和版本记录。',
    '知识片段应保存原文位置、标题层级、页码/段落、来源链接、更新时间、权限标签和向量模型版本。',
    '知识更新应支持新增、修改、删除、失效和回滚，并确保索引与原文一致。',
    '对于重复、过期、冲突、敏感或低质量知识，应进入治理流程，由知识运营或业务专家复核处理。',
    '黄金问答集应覆盖高频问题、边界问题、权限问题、无答案问题和风险问题，并随运营反馈持续更新。',
]:
    add_para(doc, item, bold_label='• ')

add_heading(doc, '9. 权限、安全与合规需求', 1)
for item in [
    '系统应遵循最小权限原则，按用户身份、组织、角色、数据标签和业务域控制知识访问。',
    '检索前应进行权限过滤，生成前应再次校验上下文片段权限，避免越权引用。',
    '系统应记录用户查询、检索结果、模型输入输出、权限过滤和管理员操作，满足审计追溯。',
    '敏感数据接入前应完成数据分级和合规评审，必要时进行脱敏、屏蔽或禁止进入模型上下文。',
    '上线前必须通过安全测试、隐私与合规检查，高危问题清零后方可上线。',
]:
    add_para(doc, item, bold_label='• ')

add_heading(doc, '10. 外部接口与集成需求', 1)
add_table(doc, ['接口/系统', '集成说明'], interfaces, widths=[4, 12])

add_heading(doc, '11. 质量与验收标准', 1)
quality_table = []
for q in qualities:
    quality_table.append([q.get('质量维度',''), q.get('指标',''), q.get('目标/门槛',''), q.get('验证方式',''), q.get('责任方','')])
add_table(doc, ['质量维度', '指标', '目标/门槛', '验证方式', '责任方'], quality_table, widths=[3, 4, 6, 4, 4])

add_heading(doc, '12. 项目阶段与里程碑要求', 1)
phase_table = []
for p in phases:
    phase_table.append([p.get('阶段',''), p.get('阶段名称',''), p.get('开始日期',''), p.get('结束日期',''), p.get('关键交付物',''), p.get('阶段退出标准','')])
add_table(doc, ['阶段', '阶段名称', '开始日期', '结束日期', '关键交付物', '阶段退出标准'], phase_table, widths=[1.5, 3, 2.4, 2.4, 5, 5])

milestone_table = []
for m in milestones:
    milestone_table.append([m.get('编号',''), m.get('里程碑',''), m.get('计划日期',''), m.get('完成判据',''), m.get('审批/责任方','')])
add_table(doc, ['编号', '里程碑', '计划日期', '完成判据', '审批/责任方'], milestone_table, widths=[1.5, 4, 2.5, 7, 4])

add_heading(doc, '13. 需求实现工作包映射', 1)
wbs_table = []
for x in wbs:
    if x.get('WBS'):
        wbs_table.append([x.get('WBS',''), x.get('过程组',''), x.get('任务名称',''), x.get('开始日期',''), x.get('结束日期',''), x.get('责任角色',''), x.get('关键交付物','')])
add_table(doc, ['WBS', '过程组', '任务名称', '开始日期', '结束日期', '责任角色', '关键交付物'], wbs_table, widths=[1.5, 2, 5, 2.3, 2.3, 3, 5])

add_heading(doc, '14. 风险与应对要求', 1)
risk_table = []
for r in risks:
    risk_table.append([r.get('风险ID',''), r.get('风险名称',''), r.get('等级',''), r.get('应对策略',''), r.get('应对措施',''), r.get('责任人','')])
add_table(doc, ['风险ID', '风险名称', '等级', '应对策略', '应对措施', '责任人'], risk_table, widths=[2, 4, 2, 3, 8, 3])

add_heading(doc, '15. 沟通与变更管理要求', 1)
comm_table = []
for c in comms:
    comm_table.append([c.get('沟通事项',''), c.get('频率',''), c.get('参与方',''), c.get('核心内容',''), c.get('负责人',''), c.get('输出物','')])
add_table(doc, ['沟通事项', '频率', '参与方', '核心内容', '负责人', '输出物'], comm_table, widths=[3, 2, 5, 6, 2.5, 4])
add_para(doc, '所有新增范围、关键指标变更、上线窗口调整、模型/供应商替换、预算变动和安全策略变更，应通过变更单提交，由 PM 组织影响评估，并经 CCB 或 Sponsor 按影响等级审批后执行。')

add_heading(doc, '16. 预算与采购约束', 1)
budget_table = []
for b in budgets:
    budget_table.append([b.get('费用类别',''), b.get('说明',''), b.get('数量/周期',''), b.get('单位',''), b.get('估算依据',''), b.get('优先级','')])
add_table(doc, ['费用类别', '说明', '数量/周期', '单位', '估算依据', '优先级'], budget_table, widths=[3, 6, 2, 2, 7, 2])
add_para(doc, '采购和选型应优先选择接口可替换、数据可迁移、具备企业级审计与 SLA 的模型、向量库和文档解析组件；MVP 阶段避免过度锁定单一供应商。')

add_heading(doc, '17. 验收与上线准入', 1)
for item in [
    'MVP Demo 通过：完成端到端问答、引用溯源、基础后台和权限入口演示。',
    'MVP 质量达标：检索命中率、引用完整性、答案有用率和时延达到项目基准。',
    '安全合规通过：高危问题清零，权限隔离、日志审计和隐私合规满足上线准入。',
    'UAT 签字通过：业务验收测试通过，遗留问题有明确豁免或解决计划。',
    'Go-Live 批准：生产发布、监控、告警、回滚和值班机制准备完成。',
    '收尾验收：验收报告签署，资产归档，运营 SOP 和运维手册完成移交。',
]:
    add_para(doc, item, bold_label='• ')

add_heading(doc, '18. 附录：需求优先级说明', 1)
add_table(doc, ['优先级', '定义'], [
    ['高', 'MVP、合规、上线准入或核心业务闭环必需，缺失将影响项目成功。'],
    ['中', '对体验、效率、扩展性或运营能力有明显价值，可根据工期分期实现。'],
    ['低', '锦上添花或后续版本优化项，不影响首期上线。'],
], widths=[3, 12])

# Footer-like final note
add_para(doc, '注：本文档为需求基线建议稿。进入详细设计前，应由 Sponsor、业务负责人、产品、架构、安全、测试和运维共同评审，并将确认后的指标写入正式需求追踪矩阵。')

# Save
if OUT.exists():
    OUT.unlink()
doc.save(OUT)
print(f'CREATED={OUT}')
print(f'SIZE={OUT.stat().st_size}')
print(f'SOURCE={SRC}')
