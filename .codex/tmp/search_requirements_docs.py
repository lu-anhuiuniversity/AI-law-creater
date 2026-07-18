from pathlib import Path
import re
from openpyxl import load_workbook
from docx import Document

root = Path(r'G:\AI-law-creater\需求文档')
keywords = ['EOS','维修','工单','企业维修','维修文档','RAG','知识库','后端','Spring','MyBatis','Elasticsearch','向量','AI']

def read_docx(p):
    doc = Document(str(p))
    parts=[]
    for para in doc.paragraphs:
        if para.text.strip(): parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            vals=[cell.text.strip().replace('\n',' ') for cell in row.cells]
            if any(vals): parts.append(' | '.join(vals))
    return '\n'.join(parts)

def read_xlsx(p):
    wb=load_workbook(str(p), data_only=True)
    parts=[]
    for ws in wb.worksheets:
        parts.append(f'# Sheet: {ws.title}')
        for row in ws.iter_rows(values_only=True):
            vals=[str(v).strip() if v is not None else '' for v in row]
            if any(vals): parts.append(' | '.join(vals))
    return '\n'.join(parts)

def read_text(p):
    for enc in ['utf-8-sig','utf-8','gbk']:
        try: return p.read_text(encoding=enc)
        except UnicodeDecodeError: pass
    return p.read_text(errors='ignore')

results=[]
outdir=Path(r'G:\AI-law-creater\.codex\tmp\requirements_text')
outdir.mkdir(parents=True, exist_ok=True)
for p in root.rglob('*'):
    if not p.is_file(): continue
    ext=p.suffix.lower()
    try:
        if ext == '.docx': text=read_docx(p)
        elif ext == '.xlsx': text=read_xlsx(p)
        elif ext in ['.md','.txt','.java','.xml','.json','.yml','.yaml']: text=read_text(p)
        else: continue
    except Exception as e:
        results.append({'file': str(p), 'error': str(e)})
        continue
    safe = re.sub(r'[\\/:*?"<>|]+','_', p.stem) + '.txt'
    (outdir/safe).write_text(text, encoding='utf-8')
    hits=[]
    lower=text.lower()
    for kw in keywords:
        count=lower.count(kw.lower())
        if count:
            snippets=[]
            for m in re.finditer(re.escape(kw), text, flags=re.I):
                start=max(0,m.start()-70); end=min(len(text),m.end()+120)
                snippets.append(text[start:end].replace('\n',' ')[:260])
                if len(snippets)>=3: break
            hits.append({'keyword':kw,'count':count,'snippets':snippets})
    results.append({'file': str(p), 'chars': len(text), 'hits': hits})

for item in results:
    print('\nFILE:', item['file'])
    if 'error' in item:
        print('ERROR:', item['error']); continue
    print('CHARS:', item['chars'])
    if not item['hits']:
        print('NO KEYWORD HITS')
    for h in item['hits']:
        print(f"- {h['keyword']}: {h['count']}")
        for s in h['snippets']:
            print('  snippet:', s)
